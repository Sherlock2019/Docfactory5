from docx import Document
import re

def extract_ordered_placeholders(docx_path):
    doc = Document(docx_path)
    placeholders = []

    def find_in_text(text):
        matches = re.findall(r"{[^}]+}", text)
        for match in matches:
            if match not in placeholders:
                placeholders.append(match)

    for para in doc.paragraphs:
        find_in_text(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                find_in_text(cell.text)

    for section in doc.sections:
        for para in section.header.paragraphs + section.footer.paragraphs:
            find_in_text(para.text)

    return placeholders
