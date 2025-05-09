from docx import Document

def fill_template(template_path, output_path, placeholder_data):
    doc = Document(template_path)

    def replace_text(text):
        for ph, content in placeholder_data.items():
            if content["type"] == "text":
                text = text.replace(ph, content["content"])
        return text

    for para in doc.paragraphs:
        para.text = replace_text(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_text(cell.text)

    for section in doc.sections:
        for para in section.header.paragraphs + section.footer.paragraphs:
            para.text = replace_text(para.text)

    doc.save(output_path)
